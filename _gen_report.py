"""Generate Snow V4 audit report as a Word document on desktop."""
import os
from pathlib import Path

desktop = Path.home() / "Desktop"
out_path = desktop / "Snow_V4_审核报告.docx"

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in (1, 2, 3):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.color.rgb = RGBColor(0x1a, 0x2b, 0x3c)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def red(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    return p

def orange(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xE0, 0x7B, 0x00)
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ══════════════════════════════════════════════
#  Cover / Title
# ══════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Snow V4 代码审核报告')
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Microsoft YaHei'
run.font.color.rgb = RGBColor(0x1a, 0x2b, 0x3c)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('Windows 桌面 AI 助手 · 全量只读审核')
run.font.size = Pt(14)
run.font.name = 'Microsoft YaHei'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026-07-17')
run.font.size = Pt(12)
run.font.name = 'Microsoft YaHei'
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════
#  1. 项目概况
# ══════════════════════════════════════════════
heading('一、项目概况', 1)

add_table(
    ['项目', '详情'],
    [
        ['语言', 'Python 3.11'],
        ['GUI 框架', 'PySide6 (Qt for Python)'],
        ['Python 源文件数', '35+'],
        ['总代码量', '约 4,500 行'],
        ['架构', 'QWebEngine 前端 + Python 后端 + LLM API'],
        ['核心功能', 'WorkerW 桌面嵌入、透明壁纸、多主题切换、Function Calling 工具调用'],
        ['AI 后端', 'DeepSeek API (deepseek-v4-flash / deepseek-v4-pro)'],
        ['TTS', 'Edge TTS (子进程隔离)'],
    ]
)

para('')
para('Snow V4 是一个运行在 Windows 桌面的本地执行型 AI 助手。它通过 PySide6 构建透明全屏覆盖窗口，使用 QWebEngine 渲染 HTML 前端，并通过 QWebChannel 实现 Python ↔ JavaScript 双向通信。核心 Agent 支持 DeepSeek 的 Function Calling 协议，具备文件读写、应用启动、网页搜索、系统命令执行等能力。')

# ══════════════════════════════════════════════
#  2. 架构概览
# ══════════════════════════════════════════════
heading('二、架构概览', 1)

para('')
heading('2.1 模块分层', 2)

add_table(
    ['层级', '模块', '职责'],
    [
        ['入口', 'main.py', '启动脚本：静默启动、崩溃捕获、日志重定向'],
        ['窗口', 'window.py', '透明全屏窗口：视频背景、WorkerW 嵌入、多模式布局、系统托盘'],
        ['桥接', 'bridge.py', 'QWebChannel 桥梁：Python ↔ JavaScript 信号/槽'],
        ['Agent', 'agent.py', 'AI Agent：流式 LLM 循环、工具调用、对话管理'],
        ['提示词', 'prompt.py + prompts/', '系统提示词构建（身份 + 行为 + 人格）'],
        ['扫描', 'scanner.py', '桌面感知：快捷方式解析（COM）、特殊桌面项检测'],
        ['分类', 'classify.py', '应用分类引擎：关键词匹配 → 六大分类'],
        ['工具', 'tools/', 'Function Calling 工具集：文件、网页、系统、桌面、应用、文档、世界面板'],
        ['服务', 'services/', '后台服务：天气、热搜、内存、TODO、TTS、工作日志、系统监控'],
        ['面板', 'panels/', 'Qt 组件：左卡片（天气/热搜）、右卡片（AI工作/待办）、世界面板'],
        ['组件', 'widgets/', '可复用组件：毛玻璃卡片、旋转Logo'],
        ['主题', 'theme_manager.py + themes/', '主题协议 V1：panel.json 配色 + HTML 模板'],
        ['设置', 'settings.py', '设置中心：模型、API密钥、TTS、主题、关于'],
        ['UI', 'ui_contract.py + ui_loader.py', 'UI 契约 + 加载器：声明接口、热切换主题'],
        ['配置', 'prompts/', '提示词规范（identity / behavior / persona）'],
    ]
)

para('')
heading('2.2 数据流', 2)
para('用户输入 → QWebEngine (HTML/JS) → QWebChannel → SnowBridge.ask() → SnowAI (LLM 流式) → token_signal → QWebEngine 渲染')
para('工具调用: SnowAI → registry.execute() → tools/*.py → 返回结果 → LLM 继续推理')
para('后台数据: QTimer 定时轮询 → bridge.get_weather/get_hotsearch → services/*.py → 面板刷新')

# ══════════════════════════════════════════════
#  3. 安全问题
# ══════════════════════════════════════════════
heading('三、安全问题 (Critical)', 1)

heading('3.1 🔴 API 密钥明文存储', 2)
red('严重度: Critical')
para('文件: config/key.txt')
para('问题: DeepSeek API Key 以明文存储在本地文件系统中，任何有文件读取权限的程序或用户都可以获取。虽然有环境变量 DEEPSEEK_API_KEY 作为回退，但本地文件存储没有任何加密或权限保护。')
para('建议: 使用 Windows DPAPI (cryptography + win32crypt) 加密存储密钥。至少应对 key.txt 设置仅当前用户可读的 ACL 权限。')

heading('3.2 🔴 SSL 证书验证被全局禁用', 2)
red('严重度: Critical')
para('文件: weather.py:21-22, _hotsearch_legacy.py:28-29, web.py:34-35')
code('ctx = ssl.create_default_context()')
code('ctx.check_hostname = False')
code('ctx.verify_mode = ssl.CERT_NONE')
para('问题: 完全绕过了 HTTPS 的中间人攻击防护。所有外部 HTTP 请求（天气、热搜、网页搜索）都暴露在 MITM 风险中，攻击者可以篡改返回数据或注入恶意内容。')
para('建议: 删除这些行。仅在调试模式下可选开启。生产代码绝不应禁用 SSL 验证。')

heading('3.3 🔴 任意命令执行 (RCE)', 2)
red('严重度: Critical')
para('文件: tools/system.py:13')
code('("run_shell", "执行PowerShell命令", {"command": ("string", "命令")}, ["command"]),')
para('问题: AI 可以通过 Function Calling 执行任意 PowerShell 命令。虽然有 30 秒超时限制，但 rm -rf、删除注册表、下载并执行恶意软件等完全可能。prompts/persona.txt 中提到"必须先征求用户同意"，但这仅依赖 LLM 遵守，没有代码级硬防护。')
para('建议: 实施命令白名单机制（如仅允许 Get-* 查询命令），或至少在执行前通过 bridge 弹窗让用户手动确认。')

heading('3.4 🔴 shell=True 命令注入风险', 2)
red('严重度: Critical')
para('文件: tools/app.py:36-37, tools/web.py:22')
code('subprocess.Popen(name, shell=True, ...)  # app.py')
code('subprocess.Popen(["cmd", "/c", "start", url], shell=True, ...)  # web.py')
para('问题: shell=True 加上未经验证的字符串输入是经典的命令注入向量。如果 AI 输出了包含 & del /f C:\\* 的字符串，会被直接执行。')
para('建议: 移除 shell=True，使用列表形式传递参数。')

# ══════════════════════════════════════════════
#  4. 高风险问题
# ══════════════════════════════════════════════
heading('四、高风险问题 (High)', 1)

heading('4.1 🟠 WorkerW 嵌入方案脆弱', 2)
orange('严重度: High')
para('文件: window.py:453-522 (_embed_to_desktop)')
para('问题: 通过 Win32 API 查找空 WorkerW 窗口并嵌入 Snow 窗口是著名的桌面壁纸替换技巧，但完全依赖 Windows 内部实现细节。EnumWindows 回调中的 nonlocal + 提前返回 False 是非标准用法。embed_epoch 竞争防护更像一个补丁。Windows 更新可能完全破坏此机制。')
para('建议: 为嵌入失败提供优雅降级（作为普通置底窗口），而非停留在半嵌入的不可用状态。')

heading('4.2 🟠 重启助手不可靠', 2)
orange('严重度: High')
para('文件: _restart_helper.py')
code('time.sleep(3.0)  # 给旧进程足够时间退出')
para('问题: 固定 3 秒睡眠是不稳健的。如果旧进程需要更长时间退出（等待 TTS 播放、COM 释放），新进程会启动失败或产生资源冲突。')
para('建议: 改为轮询检查旧进程 PID 是否还存在 (psutil.process_iter)，确认完全退出后再启动新进程。')

heading('4.3 🟠 COM 初始化模式有泄漏风险', 2)
orange('严重度: High')
para('文件: scanner.py:12-27, tools/world.py:78-83')
para('问题: 如果 Dispatch 或 CreateShortCut 抛出异常，CoUninitialize() 不会执行，导致 COM 资源泄漏。虽然有嵌套 try/except 补救，但代码逻辑混乱。')
para('建议: 使用 with 上下文管理器或在 finally 块中统一清理 COM 资源。')

heading('4.4 🟠 应用分类引擎不可靠', 2)
orange('严重度: High')
para('文件: classify.py')
para('问题: 基于关键词匹配的分类方式有根本性缺陷。例如"微信"同时出现在"社交"、"此刻"和"办公"三个分类中，导致分类结果不确定。英文命名的应用容易被误分 ("code" 会匹配 VS Code 但也会误匹配其他含 "code" 字样的程序)。')
para('建议: 引入分类优先级或置信度机制。利用已有的 app_cache.json 持久化用户手动调整的分类结果。')

# ══════════════════════════════════════════════
#  5. 中等问题
# ══════════════════════════════════════════════
heading('五、中等问题 (Medium)', 1)

heading('5.1 🟡 缺少任何形式的测试', 2)
para('整个项目没有单元测试、集成测试或端到端测试。考虑到涉及的复杂性（多线程、COM、WorkerW、TTS 子进程、LLM 流式响应），这很危险。')
para('建议: 至少为核心模块（classify、services/todos、tools/registry）添加 pytest 单元测试。')

heading('5.2 🟡 window.py 过长 (1120+ 行)', 2)
para('单文件同时管理视频播放、WorkerW 嵌入、UI 主题切换、多模式布局（chat_only/transparent/morandi/红丝带）、系统托盘、桌面图标显隐、渲染崩溃恢复。职责过多。')
para('建议: 拆分为 VideoManager、DesktopEmbedder、LayoutManager 等独立模块。')

heading('5.3 🟡 全局单例模式破坏模块边界', 2)
para('多处使用 from snow.bridge import _bridge_instance 获取全局桥接实例。tools/snow.py、tools/world.py、settings.py 等跨层依赖这个全局变量，使单元测试和多实例化几乎不可能。')
para('建议: 使用依赖注入或将 bridge 实例作为参数传入工具函数。')

heading('5.4 🟡 TTS 临时文件清理不可靠', 2)
para('播放完成后 os.remove(wav) 在 finally 块中执行，但如果进程被强制终止，临时文件会残留。_sweep_temp_audio() 只在引擎初始化时运行一次，不会定期清理。')
para('建议: 使用 tempfile.NamedTemporaryFile(delete=True) 或 Windows 临时目录的自动清理。添加定期清理任务。')

heading('5.5 🟡 文件读取无大小限制', 2)
para('tools/file.py 的 read_file 先将整个文件读入内存再截断到 4000 字符。如果 AI 尝试读取大文件，会导致内存溢出。')
para('建议: 先用 os.path.getsize() 检查文件大小，超过阈值（如 10MB）直接拒绝。')

heading('5.6 🟡 分类缓存无版本管理', 2)
para('tools/classify.py 的 app_cache.json 没有版本号或 schema 验证。如果日后修改分类逻辑，旧缓存数据会导致不一致。')

# ══════════════════════════════════════════════
#  6. 代码质量问题
# ══════════════════════════════════════════════
heading('六、代码质量问题 (Low)', 1)

bullets = [
    '缺少类型标注 — 整个项目没有任何 Python type hints。对于多线程桌面应用，类型标注可有效防止运行时错误。',
    '硬编码魔法数字 — agent.py 中的 300（历史）、100（最大轮次）、16384（max_tokens）、90（超时），window.py 中的 15000（刷新间隔）、2000（嵌入延迟）等散布各处，应提取为配置文件或模块常量。',
    '日志系统原始 — 所有模块通过 open(path, "a").write() 手动写日志，无统一日志级别、格式、轮转机制。多个线程竞争写入同一个 diag.log 缺乏线程安全。建议使用 logging 模块。',
    'world.py 文档字符串乱码 — 第一行包含 "閳" 编码错误产物。',
    '中英文混杂 — 注释、文档字符串、UI 文本时而中文时而英文，缺乏一致性。',
    '空的 __init__.py — 包初始化文件均为空，应包含公共 API 导入。',
]
for b in bullets:
    bullet(b)

# ══════════════════════════════════════════════
#  7. 架构评价
# ══════════════════════════════════════════════
heading('七、架构评价', 1)

heading('7.1 优点 ✅', 2)
pros = [
    '模块化良好 — services/、tools/、panels/、widgets/ 分层清晰，职责边界明确。',
    '工具注册表可扩展 — registry.py 的自动扫描模式使添加新工具极其简单，只需在 tools/ 下新建文件并定义 TOOLS + HANDLERS。',
    '主题系统设计优秀 — ui_contract.py（契约）+ ui_loader.py（加载）+ theme_manager.py（配色）三层设计让 UI 和 Qt 配色完全独立，panel.json 可自由定制。',
    '多模式共存 — chat_only / morandi / 透明壁纸 / 视频背景四种主题模式共享同一后端，不需要多套代码。',
    '防崩溃措施 — faulthandler 捕获 C 层崩溃、线程异常钩子记录未捕获异常、渲染进程崩溃自动恢复。',
    'Function Calling 流式处理 — agent.py 的 SSE 解析 + tool_calls 累积方案正确处理了 DeepSeek API 的分块特性（多个 delta 拼接 tool_calls）。',
]
for p in pros:
    bullet(p)

heading('7.2 缺点 ❌', 2)
cons = [
    '嵌入桌面方案脆弱 — 依赖 Windows 未公开行为，升级风险高。',
    '全局可变状态 — _bridge_instance 跨层引用难以测试和维护。',
    '无测试覆盖 — 迭代重构风险高。',
    '安全性不足 — SSL 禁用、明文密钥、任意命令执行。',
]
for c in cons:
    bullet(c)

# ══════════════════════════════════════════════
#  8. 安全风险汇总
# ══════════════════════════════════════════════
heading('八、安全风险汇总', 1)

add_table(
    ['风险', '严重度', '文件'],
    [
        ['API Key 明文存储', '🔴 Critical', 'config/key.txt'],
        ['SSL 验证禁用 (MITM 攻击)', '🔴 Critical', 'weather.py, _hotsearch_legacy.py, web.py'],
        ['任意 PowerShell 执行', '🔴 Critical', 'tools/system.py'],
        ['shell=True 命令注入', '🔴 Critical', 'tools/app.py, tools/web.py'],
        ['无用户确认的破坏性操作', '🟠 High', '多个工具模块'],
        ['明文密钥可被任意程序读取', '🟠 High', 'config/key.txt'],
    ]
)

# ══════════════════════════════════════════════
#  9. 总结与评分
# ══════════════════════════════════════════════
heading('九、总结与评分', 1)

para('Snow V4 是一个功能丰富、架构清晰的 Windows 桌面 AI 助手。如果这是一个个人项目或原型，代码质量已经相当不错。它展示了桌面 AI 助手的完整愿景：桌面感知、应用控制、文件操作、网页搜索等能力一应俱全。')

para('但如果计划公开分发或长期维护，必须重点解决以下三个方面：')

bullet('安全加固 — SSL 验证恢复、密钥加密存储、命令执行白名单、用户二次确认')
bullet('稳定性 — WorkerW 嵌入回退机制、COM 资源正确清理、重启进程监控')
bullet('可维护性 — 拆分 window.py、添加单元测试、统一日志系统')

para('')
heading('评分', 2)

add_table(
    ['维度', '评分 (满分10)', '说明'],
    [
        ['功能完整度', '8/10', 'AI 对话、工具调用、主题切换、桌面集成均完善'],
        ['代码质量', '6/10', '分层清晰但缺少类型标注、测试和统一日志'],
        ['安全性', '4/10', '明文密钥、SSL 禁用、任意命令执行是硬伤'],
        ['可维护性', '5/10', '模块化好但 window.py 过长，全局状态难测试'],
        ['可扩展性', '7/10', '工具注册表 + 主题协议设计优秀，添加功能便捷'],
        ['健壮性', '5/10', '防崩溃措施到位但 COM/嵌入方案脆弱'],
        ['综合评分', '5.8/10', '功能先行，安全和工程化跟进不足'],
    ]
)
para('')

# ── File list ──
heading('十、审核文件清单', 1)

files = [
    'main.py',
    '_restart_helper.py',
    'restore_icons.py',
    'requirements.txt',
    'snow/__init__.py (空)',
    'snow/agent.py',
    'snow/bridge.py',
    'snow/classify.py',
    'snow/prompt.py',
    'snow/scanner.py',
    'snow/settings.py',
    'snow/theme_manager.py',
    'snow/ui_contract.py',
    'snow/ui_loader.py',
    'snow/window.py',
    'snow/panels/__init__.py (空)',
    'snow/panels/left_card.py',
    'snow/panels/right_card.py',
    'snow/panels/world.py',
    'snow/services/__init__.py (空)',
    'snow/services/_edge_tts_worker.py',
    'snow/services/_hotsearch_legacy.py',
    'snow/services/_tts_engine.py',
    'snow/services/hotsearch.py',
    'snow/services/memory.py',
    'snow/services/monitor.py',
    'snow/services/todos.py',
    'snow/services/tts.py',
    'snow/services/weather.py',
    'snow/services/worklog.py',
    'snow/tools/__init__.py (空)',
    'snow/tools/app.py',
    'snow/tools/classify.py',
    'snow/tools/desktop.py',
    'snow/tools/document.py',
    'snow/tools/file.py',
    'snow/tools/registry.py',
    'snow/tools/snow.py',
    'snow/tools/system.py',
    'snow/tools/web.py',
    'snow/tools/world.py',
    'snow/widgets/__init__.py (空)',
    'snow/widgets/glass_card.py',
    'snow/widgets/logo.py',
    'config/model_config.json',
    'config/ui.json',
    'config/key.txt',
    'prompts/identity.txt',
    'prompts/behavior.txt',
    'prompts/persona.txt',
    'themes/installed/chat-only/manifest.json',
    'themes/installed/morandi/manifest.json',
    'themes/template/manifest.json',
]
for f in files:
    bullet(f)

# ── Save ──
doc.save(str(out_path))
print(f'Report saved to: {out_path}')
