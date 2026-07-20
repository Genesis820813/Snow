"""Tool: document operations — write_word, write_excel, find_documents."""
import os as _os
from pathlib import Path

TOOLS = [
    ("find_documents", "\u641c\u7d22\u6587\u6863", {"keyword": ("string", "\u5173\u952e\u8bcd"), "filetype": ("string", "\u6269\u5c55\u540d")}, []),
    ("write_word", "\u5199\u5165Word", {"path": ("string", "\u8def\u5f84"), "content": ("string", "\u5185\u5bb9"), "mode": ("string", "create|append")}, ["path", "content"]),
    ("write_excel", "\u5199\u5165Excel", {"path": ("string", "\u8def\u5f84"), "content": ("string", "CSV\u6570\u636e"), "mode": ("string", "create|append")}, ["path", "content"]),
]


def _find_documents(args):
    keyword = args.get("keyword", "")
    filetype = args.get("filetype", "")
    dirs = [Path.home() / "Desktop", Path.home() / "Documents"]
    exts = [filetype.lower()] if filetype else [".docx", ".xlsx", ".pdf", ".txt", ".doc", ".xls", ".ppt", ".pptx", ".md", ".csv"]
    results = []
    for d in dirs:
        if not d.exists():
            continue
        for root, dirs2, files in _os.walk(str(d)):
            dirs2[:] = [x for x in dirs2 if not x.startswith(".")]
            for f in files:
                if f.startswith(("~", ".")):
                    continue
                fp = _os.path.join(root, f)
                if any(f.lower().endswith(e) for e in exts):
                    if not keyword or keyword.lower() in f.lower():
                        results.append(fp)
                        if len(results) >= 30:
                            break
            if len(results) >= 30:
                break
    if not results:
        return "\u672a\u627e\u5230\u5339\u914d\u6587\u6863"
    return "\n".join(results)


def _write_word(args):
    path = args.get("path", "")
    content = args.get("content", "")
    mode = args.get("mode", "create")
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
        doc = Document() if (mode == "create" or not p.exists()) else Document(str(p))
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        doc.save(str(p))
        return f"\u5df2\u5199\u5165 Word: {p.name}"
    except ImportError:
        tp = p.with_suffix(".txt")
        tp.write_text(content, "utf-8")
        return f"\u5df2\u5199\u5165\u7eaf\u6587\u672c: {tp.name}"


def _write_excel(args):
    path = args.get("path", "")
    content = args.get("content", "")
    mode = args.get("mode", "create")
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    try:
        from openpyxl import Workbook, load_workbook
        if mode == "create" or not p.exists():
            wb = Workbook(); ws = wb.active
        else:
            wb = load_workbook(str(p))
            ws = wb.create_sheet() if mode == "append" else wb.active
        for ri, line in enumerate(content.split("\n"), 1):
            if line.strip():
                for ci, cell in enumerate(line.split(","), 1):
                    ws.cell(row=ri, column=ci, value=cell.strip())
        wb.save(str(p))
        return f"\u5df2\u5199\u5165 Excel: {p.name}"
    except ImportError:
        cp = p.with_suffix(".csv")
        cp.write_text(content, "utf-8")
        return f"\u5df2\u5199\u5165 CSV: {cp.name}"


HANDLERS = {
    "find_documents": _find_documents,
    "write_word": _write_word,
    "write_excel": _write_excel,
}
